import os
import json
import uuid
import random
import logging
import spacy
from docx import Document
from typing import List, Dict, Any, Optional
import asyncio
import aiohttp
from collections import defaultdict
from async_lru import alru_cache
from scholarly import scholarly, ProxyGenerator
import time
from asyncio import Semaphore, Queue
import weakref
from dataclasses import dataclass
from concurrent.futures import ThreadPoolExecutor
from app.core.gemini_helper import get_document_context_with_gemini


logging.basicConfig(level=logging.INFO)

@dataclass
class SearchResult:
    title: str
    authors: List[str]
    year: Optional[int]
    venue: Optional[str]
    url: Optional[str]
    citations: int
    source: str
    relevance_score: float = 0.0
    context_match: float = 0.0

class CircuitBreaker:
    def __init__(self, failure_threshold: int = 5, recovery_timeout: int = 60):
        self.failure_threshold = failure_threshold
        self.recovery_timeout = recovery_timeout
        self.failure_count = 0
        self.last_failure_time = None
        self.state = 'closed'
    
    async def call(self, func):
        if self.state == 'open':
            if time.time() - self.last_failure_time > self.recovery_timeout:
                self.state = 'half-open'
            else:
                raise Exception("Circuit breaker is open")
        
        try:
            result = await func()
            if self.state == 'half-open':
                self.state = 'closed'
                self.failure_count = 0
            return result
        except Exception as e:
            self.failure_count += 1
            self.last_failure_time = time.time()
            if self.failure_count >= self.failure_threshold:
                self.state = 'open'
            raise e

class AcademicCitationProcessor:
    def __init__(self, style="APA", search_providers=None, threshold=0.0, top_k=5, max_api_calls=None, max_concurrent=50):
        self.style = style
        self.search_providers = search_providers or ["semantic_scholar", "crossref", "openalex"]
        self.threshold = threshold
        self.top_k = top_k
        self.max_api_calls = max_api_calls
        self.max_concurrent = max_concurrent
        self.api_call_count = 0
        self.matched_paper_titles = []
        
        self.research_context = ""
        self.document_category = ""
        self.field_keywords = []
        
        self.semaphore = Semaphore(max_concurrent)
        self.session_cache = weakref.WeakValueDictionary()
        self.circuit_breakers = {provider: CircuitBreaker() for provider in self.search_providers}
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        try:
            self.nlp = spacy.load("en_core_web_sm")
            self.nlp.max_length = 2000000
        except OSError:
            logging.error("SpaCy model 'en_core_web_sm' not found. Please install it with: python -m spacy download en_core_web_sm")
            raise

    def enhance_query_with_context(self, original_query: str, sentence_context: str = "") -> str:
        enhanced_query = original_query
        
        if self.research_context:
            enhanced_query = f"{self.research_context} {enhanced_query}"
        
        if self.document_category:
            enhanced_query = f"{enhanced_query} {self.document_category}"
        
        if self.field_keywords:
            field_terms = " ".join(self.field_keywords[:3])
            enhanced_query = f"{enhanced_query} {field_terms}"
        
        return self.clean_query(enhanced_query)

    def _calculate_api_limits_and_eta(self, total_sentences: int) -> tuple:
        if self.max_api_calls is not None:
            return self.max_api_calls, 0
        
        citation_rate = 1
        avg_providers_per_search = 5
        estimated_citations = int(total_sentences * citation_rate)
        calculated_max_calls = min(estimated_citations * avg_providers_per_search, 150)
        calculated_max_calls = 1000
        
        avg_time_per_call = 0.3
        estimated_eta_seconds = calculated_max_calls * avg_time_per_call
        
        self.max_api_calls = int(calculated_max_calls)
        
        return self.max_api_calls, estimated_eta_seconds

    def smart_sentence_selection(self, all_sentences: list, max_sentences: int = None) -> list:
        if not all_sentences:
            return []
        
        if max_sentences is None:
            max_sentences = min(len(all_sentences), 500)
        
        if len(all_sentences) <= max_sentences:
            return all_sentences
        
        academic_keywords = {
            'study', 'research', 'analysis', 'data', 'results', 'findings', 'evidence',
            'method', 'approach', 'theory', 'model', 'framework', 'hypothesis',
            'significant', 'correlation', 'impact', 'effect', 'relationship',
            'according', 'reported', 'demonstrated', 'showed', 'indicated',
            'suggests', 'concluded', 'observed', 'measured', 'tested',
            'literature', 'review', 'quantitative', 'qualitative', 'experiment',
            'survey', 'interview', 'protocol', 'algorithm', 'simulation'
        }
        
        context_keywords = set(self.field_keywords + [self.research_context, self.document_category])
        context_keywords = {kw.lower() for kw in context_keywords if kw}
        
        scored_sentences = []
        for s in all_sentences:
            text_lower = s['text'].lower()
            score = sum(1 for kw in academic_keywords if kw in text_lower)
            context_score = sum(1 for kw in context_keywords if kw in text_lower)
            
            if len(s['text'].split()) > 10:
                score += 2
            
            score += context_score * 2
            
            scored_sentences.append((s, score))
        
        scored_sentences.sort(key=lambda x: x[1], reverse=True)
        return [s[0] for s in scored_sentences[:max_sentences]]

    def is_dynamic_heading(self, para) -> bool:
        text = para.text.strip()
        if not text:
            return False
        try:
            if "heading" in para.style.name.lower():
                return True
        except Exception:
            pass
        words = text.split()
        return len(words) < 8 and not any(punct in text for punct in ".?!;:")

    def clean_query(self, query: str) -> str:
        query = query.strip()
        if query.startswith('-') or query.startswith('•'):
            query = query[1:].strip()
        if query and query[0].isdigit() and '.' in query[:5]:
            query = query.split('.', 1)[1].strip()
        words = query.split()
        return ' '.join(words[:15])

    async def get_session(self) -> aiohttp.ClientSession:
        session_id = id(asyncio.current_task())
        if session_id not in self.session_cache:
            connector = aiohttp.TCPConnector(
                limit=100,
                limit_per_host=20,
                ttl_dns_cache=300,
                use_dns_cache=True,
                keepalive_timeout=60,
                enable_cleanup_closed=True
            )
            timeout = aiohttp.ClientTimeout(total=8, connect=3)
            session = aiohttp.ClientSession(
                connector=connector, 
                timeout=timeout,
                headers={'User-Agent': 'Academic Citation Processor 1.0'}
            )
            self.session_cache[session_id] = session
        return self.session_cache[session_id]

    @alru_cache(maxsize=2048)
    async def search_all_providers_async(self, query: str, context: str = "", max_results: int = None) -> List[SearchResult]:
        async with self.semaphore:
            if self.api_call_count >= self.max_api_calls:
                return []
            
            enhanced_query = self.enhance_query_with_context(query, context)
            
            if not enhanced_query or len(enhanced_query) < 5:
                return []
            
            max_results = max_results or self.top_k
            session = await self.get_session()
            
            tasks = []
            for provider in self.search_providers:
                if self.api_call_count >= self.max_api_calls:
                    break
                
                self.api_call_count += 1
                task = self._search_provider_with_circuit_breaker(session, provider, enhanced_query, max_results)
                tasks.append(task)
            
            if not tasks:
                return []
            
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            all_papers = []
            seen_titles = set()
            
            for result in results:
                if isinstance(result, Exception):
                    continue
                
                for paper in result:
                    title = paper.title.lower().strip() if paper.title else ''
                    if title and title not in seen_titles:
                        seen_titles.add(title)
                        paper.context_match = self.calculate_context_match(paper, context)
                        all_papers.append(paper)
            
            return all_papers

    def calculate_context_match(self, paper: SearchResult, context: str) -> float:
        context_score = 0.0
        
        paper_text = f"{paper.title} {paper.venue or ''}".lower()
        context_lower = context.lower()
        
        if self.research_context:
            research_words = set(self.research_context.lower().split())
            paper_words = set(paper_text.split())
            overlap = len(research_words.intersection(paper_words))
            context_score += (overlap / len(research_words)) * 0.4 if research_words else 0
        
        if self.field_keywords:
            field_matches = sum(1 for kw in self.field_keywords if kw.lower() in paper_text)
            context_score += (field_matches / len(self.field_keywords)) * 0.3
        
        if self.document_category and self.document_category.lower() in paper_text:
            context_score += 0.3
        
        return min(context_score, 1.0)

    async def _search_provider_with_circuit_breaker(self, session: aiohttp.ClientSession, provider: str, query: str, max_results: int) -> List[SearchResult]:
        try:
            circuit_breaker = self.circuit_breakers[provider]
            return await circuit_breaker.call(lambda: self._search_provider_async(session, provider, query, max_results))
        except Exception:
            return []

    async def _search_provider_async(self, session: aiohttp.ClientSession, provider: str, query: str, max_results: int) -> List[SearchResult]:
        try:
            if provider == 'semantic_scholar':
                return await self._search_semantic_scholar_async(session, query, max_results)
            elif provider == 'crossref':
                return await self._search_crossref_async(session, query, max_results)
            elif provider == 'openalex':
                return await self._search_openalex_async(session, query, max_results)
            elif provider == 'google_scholar':
                return await self._search_google_scholar_async(query, max_results)
        except Exception as e:
            logging.debug(f"Failed to search {provider}: {e}")
            return []

    async def _search_semantic_scholar_async(self, session: aiohttp.ClientSession, query: str, max_results: int) -> List[SearchResult]:
        url = 'https://api.semanticscholar.org/graph/v1/paper/search'
        params = {
            'query': query, 
            'limit': max_results, 
            'fields': 'title,authors,year,venue,citationCount,url,publicationDate,fieldsOfStudy'
        }
        
        async with session.get(url, params=params) as response:
            if response.status == 429:
                await asyncio.sleep(0.5)
                return []
            response.raise_for_status()
            data = await response.json()
            
            results = []
            for p in data.get('data', []):
                if p.get('title') and p.get('authors'):
                    results.append(SearchResult(
                        title=p.get('title'),
                        authors=[a.get('name') for a in p.get('authors', [])],
                        year=p.get('year'),
                        venue=p.get('venue'),
                        url=p.get('url'),
                        citations=p.get('citationCount', 0),
                        source='Semantic Scholar'
                    ))
            return results

    async def _search_crossref_async(self, session: aiohttp.ClientSession, query: str, max_results: int) -> List[SearchResult]:
        url = 'https://api.crossref.org/works'
        params = {'query': query, 'rows': max_results, 'sort': 'relevance'}
        
        async with session.get(url, params=params) as response:
            response.raise_for_status()
            data = await response.json()
            
            results = []
            for item in data.get('message', {}).get('items', []):
                authors = [f"{a.get('given', '')} {a.get('family', '')}".strip() for a in item.get('author', [])]
                if not item.get('title') or not authors:
                    continue
                
                year = None
                if item.get('published-print', {}).get('date-parts'):
                    year = item.get('published-print', {}).get('date-parts', [[None]])[0][0]
                
                results.append(SearchResult(
                    title=' '.join(item.get('title', [])),
                    authors=authors,
                    year=year,
                    venue=' '.join(item.get('container-title', [])),
                    url=item.get('URL'),
                    citations=item.get('is-referenced-by-count', 0),
                    source='Crossref'
                ))
            return results

    async def _search_openalex_async(self, session: aiohttp.ClientSession, query: str, max_results: int) -> List[SearchResult]:
        url = 'https://api.openalex.org/works'
        params = {'search': query, 'per-page': max_results, 'sort': 'relevance_score:desc'}
        
        async with session.get(url, params=params) as response:
            response.raise_for_status()
            data = await response.json()
            
            results = []
            for work in data.get('results', []):
                authors = [a['author'].get('display_name') for a in work.get('authorships', [])]
                if not work.get('title') or not authors:
                    continue
                
                venue = work.get('primary_location', {}).get('source', {}).get('display_name')
                
                results.append(SearchResult(
                    title=work.get('title'),
                    authors=authors,
                    year=work.get('publication_year'),
                    venue=venue,
                    url=work.get('primary_location', {}).get('landing_page_url'),
                    citations=work.get('cited_by_count', 0),
                    source='OpenAlex'
                ))
            return results

    async def _search_google_scholar_async(self, query: str, max_results: int) -> List[SearchResult]:
        try:
            loop = asyncio.get_running_loop()
            search_results = await loop.run_in_executor(
                self.executor, 
                lambda: list(scholarly.search_pubs(query))
            )
            
            results = []
            for i, pub in enumerate(search_results):
                if i >= max_results:
                    break
                
                bib = pub.get('bib', {})
                results.append(SearchResult(
                    title=bib.get('title'),
                    authors=bib.get('author', []) if isinstance(bib.get('author'), list) else [bib.get('author')] if bib.get('author') else [],
                    year=int(bib.get('pub_year')) if bib.get('pub_year') and str(bib.get('pub_year')).isdigit() else None,
                    venue=bib.get('venue'),
                    url=pub.get('pub_url'),
                    citations=pub.get('num_citations', 0),
                    source='Google Scholar'
                ))
            return results
        except Exception as e:
            logging.debug(f"Error searching Google Scholar: {e}")
            return []

    def calculate_relevance_score(self, sentence: str, paper: SearchResult) -> float:
        if not sentence or not isinstance(sentence, str) or not paper.authors:
            return 0.0
        
        sentence_lower = sentence.lower()
        title_lower = paper.title.lower() if paper.title else ''
        
        stop_words = {'the', 'a', 'an', 'and', 'or', 'in', 'on', 'to', 'for', 'of', 'is', 'are', 'was', 'were', 'with', 'by', 'from', 'this', 'that'}
        sentence_words = set(sentence_lower.split()) - stop_words
        title_words = set(title_lower.split()) - stop_words
        
        if not sentence_words:
            return 0.0
        
        title_overlap = len(sentence_words.intersection(title_words)) / len(sentence_words)
        score = title_overlap * 0.5
        
        score += paper.context_match * 0.3
        
        if paper.year and paper.year >= 2020:
            score *= 1.3
        elif paper.year and paper.year >= 2015:
            score *= 1.15
        elif paper.year and paper.year >= 2010:
            score *= 1.05
        
        if paper.citations > 200:
            score *= 1.2
        elif paper.citations > 100:
            score *= 1.15
        elif paper.citations > 50:
            score *= 1.1
        
        return min(score, 1.0)

    async def batch_process_sentences_async(self, sentences: list) -> list:
        semaphore = Semaphore(self.max_concurrent)
        
        async def process_with_semaphore(sentence_data):
            async with semaphore:
                return await self.process_single_sentence_async(sentence_data)
        
        tasks = [process_with_semaphore(s) for s in sentences if self.api_call_count < self.max_api_calls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        citations = []
        for res in results:
            if isinstance(res, Exception):
                logging.debug(f"Error processing sentence: {res}")
            elif res:
                citations.append(res)
        
        return citations

    async def process_single_sentence_async(self, sentence_data: dict) -> dict:
        try:
            sentence_text = sentence_data['text']
            papers = await self.search_all_providers_async(sentence_text, sentence_text)
            
            if not papers:
                return None
            
            for paper in papers:
                paper.relevance_score = self.calculate_relevance_score(sentence_text, paper)
            
            relevant_papers = [p for p in papers if p.relevance_score >= self.threshold]
            
            if not relevant_papers:
                return None
            
            best_paper = max(relevant_papers, key=lambda x: x.relevance_score + x.context_match)
            
            if not best_paper.year or (best_paper.year and best_paper.year < 2015):
                return None

            return {
                "id": str(uuid.uuid4()),
                "original_sentence": sentence_text,
                "paper_details": {
                    "title": best_paper.title,
                    "authors": best_paper.authors,
                    "year": str(best_paper.year) if best_paper.year else "n.d.",
                    "url": best_paper.url or '',
                    "venue": best_paper.venue or '',
                    "citations": best_paper.citations,
                    "relevance_score": round(best_paper.relevance_score, 3),
                    "context_match": round(best_paper.context_match, 3),
                    "source": best_paper.source,
                },
                "metadata": {
                    "paragraph_index": sentence_data['actual_para_idx'],
                    "sentence_index": sentence_data['sent_idx'],
                }
            }
        except Exception as e:
            logging.debug(f"Error in process_single_sentence_async: {e}")
            return None

    async def prepare_citations_for_review(self, input_path: str, max_paragraphs: int = 1000) -> Dict[str, Any]:
        start_time = time.time()
        logging.info(f"Starting contextual citation processing for: '{input_path}'")
        
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file '{input_path}' does not exist.")

        doc = Document(input_path)
        
        # Extract full document content for Gemini
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Get context from Gemini
        context_data = await get_document_context_with_gemini(full_text)
        self.research_context = context_data.get("research_context", "")
        self.document_category = context_data.get("document_category", "")
        self.field_keywords = context_data.get("field_keywords", [])
        
        logging.info(f"Gemini context acquired: Category='{self.document_category}', Context='{self.research_context}'")

        paragraphs_to_process = doc.paragraphs[:min(len(doc.paragraphs), max_paragraphs)]
        
        all_sentences = []
        for para_idx, para in enumerate(paragraphs_to_process):
            if not para.text.strip() or self.is_dynamic_heading(para):
                continue
            
            try:
                sentences = list(self.nlp(para.text.strip()).sents)
                for sent_idx, sent in enumerate(sentences, 1):
                    text = sent.text.strip()
                    if len(text) >= 15 and len(text.split()) >= 5:
                        all_sentences.append({
                            'text': text,
                            'actual_para_idx': para_idx + 1,
                            'sent_idx': sent_idx
                        })
            except Exception as e:
                logging.debug(f"Error tokenizing paragraph {para_idx}: {e}")

        total_sentences = len(all_sentences)
        calculated_max_calls, estimated_eta = self._calculate_api_limits_and_eta(total_sentences)
        selected_sentences = self.smart_sentence_selection(all_sentences, min(total_sentences, 500))
        
        logging.info(f"Processing {len(selected_sentences)} sentences with context: {self.research_context}")
        
        citations = await self.batch_process_sentences_async(selected_sentences)
        
        processing_time = time.time() - start_time
        logging.info(f"Contextual citation processing completed in {processing_time:.2f} seconds")
        
        return {
            "document_id": str(uuid.uuid4()),
            "total_citations": len(citations),
            "citations": citations,
            "context_info": {
                "research_context": self.research_context,
                "document_category": self.document_category,
                "field_keywords": self.field_keywords,
                "detected_domain": self.document_category # Use the Gemini-detected category
            },
            "diagnostics": {
                "processed_paragraphs": len(set(s['actual_para_idx'] for s in selected_sentences)),
                "processed_sentences": len(selected_sentences),
                "api_calls_made": self.api_call_count,
                "max_api_calls": calculated_max_calls,
                "processing_time_seconds": round(processing_time, 2),
                "sentences_per_second": round(len(selected_sentences) / processing_time, 2),
            }
        }

    async def cleanup(self):
        for session in self.session_cache.values():
            if not session.closed:
                await session.close()
        self.executor.shutdown(wait=True)