import os
import logging
from docx import Document
from typing import List, Optional
from dataclasses import dataclass
import spacy
import numpy as np
from pymongo import MongoClient
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from datapipeline.core.mongo_client import MongoDBVectorStoreManager, MongoDBAtlasVectorSearch
from datapipeline.models.papers import Papers
from datapipeline.core.database import get_session_with_ctx_manager
from datapipeline.core.utils import embeddings
from datapipeline.core.constants import MONGODB_ATLAS_CLUSTER_URI, MONGO_DB_NAME
from app.core.headings import headers

# Initialize SpaCy and logging
nlp = spacy.load("en_core_web_sm")
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

@dataclass
class PaperMetadata:
    title: str
    authors: List[str]
    year: str
    category: str

class SemanticCitationInserter:
    def __init__(
        self, 
        mongo_uri: str, 
        db_name: str, 
        collection_name: str,
        threshold: float = 0.75, 
        top_k: int = 3,
        headers = headers
    ):
        self.threshold = threshold
        self.top_k = top_k
        self.collection_name = collection_name
        self.embedding_model = embeddings
        
        # Initialize MongoDB manager
        self.manager = MongoDBVectorStoreManager(
            connection_string=mongo_uri,
            db_name=db_name
        )

        # Initialize headings of various sections
        # TODO: in future work on an implementation that identifies them correctly.
        self.headers = headers


    def is_heading(self, sentence: str) -> bool:
        """
        Checks if a sentence matches any of the predefined headings.
        """
        return sentence.strip() in self.headers


    def find_relevant_papers(self, sentence_query: str):
        """
        Retrieve relevant papers using the retriever.

        Parameters:
        - sentence_query (str): The query sentence to find relevant papers.

        Returns:
        - List[Document]: A list of relevant documents.
        """
        logging.info("Searching for relevant papers using retriever.")
        retriever = self.manager.get_retriever(
            self.collection_name,
            search_type="similarity",
            k=self.top_k
        )
        results = retriever.get_relevant_documents(sentence_query)
        logging.info(f"Found {len(results)} relevant papers.")
        return results
    

    def retrieve_full_text(self, paper_metadata: dict) -> Optional[str]:
        """"
        Retrieve the full text of a paper from metadata stored in MongoDB.

        Parameters:
        - paper_metadata (dict): The metadata of the paper.

        Returns:
        - Optional[str]: The full text content of the paper.
        """
        return paper_metadata.get("content")
    

    def filter_relevant_collection(self) -> MongoDBAtlasVectorSearch:
        """
        Filters the MongoDB collection to include only relevant papers.
        """
        collection = self.manager.get_or_create_collection(self.collection_name)
        vector_store = MongoDBAtlasVectorSearch(
            collection=collection,
            embedding=self.embedding_model,
            index_name=f"{self.collection_name}_index",
            relevance_score_fn="cosine"
        )
        filtered_collection = [
            paper for paper in collection.find({"metadata.title": {"$in": self.relevant_papers}})
        ]
        vector_store.add_documents([
            LangChainDocument(
                page_content=doc.get("content", ""),
                metadata=doc.get("metadata", {})
            )
            for doc in filtered_collection
        ])
        return vector_store
    

    # def process_document(self, input_doc_path: str, output_doc_path: str):
    #     logging.info(f"Processing document: {input_doc_path}")
    #     doc = Document(input_doc_path)

    #     for paragraph_index, para in enumerate(doc.paragraphs, start=1):
    #         if not para.text.strip():
    #             continue

    #         logging.info(f"Processing paragraph {paragraph_index}.")
    #         sentences = list(nlp(para.text).sents)
    #         new_para_text = ""

    #         for sentence_index, sent in enumerate(sentences, start=1):
    #             sentence_text = sent.text.strip()
    #             if not sentence_text:
    #                 continue

    #             # Check if the sentence is a standalone heading
    #             if self.is_heading(sentence_text):
    #                 logging.info(f"Skipping heading: {sentence_text}")
    #                 new_para_text += f" {sentence_text}"  # Preserve headings in the output
    #                 continue

    #             logging.info(f"Processing sentence {sentence_index}: {sentence_text}")
    #             relevant_papers = self.find_relevant_papers(sentence_text)

    #             if relevant_papers:
    #                 citation_texts = []

    #                 for paper_doc in relevant_papers:
    #                     paper_metadata = paper_doc.metadata
    #                     title = paper_metadata.get("title")

    #                     if title:
    #                         full_text = self.retrieve_full_text(paper_metadata)
    #                         if full_text:
    #                             citation_texts.append(f"({title})")

    #                 if citation_texts:
    #                     combined_citation = " ".join(citation_texts)
    #                     sentence_text += f" {combined_citation}"

    #             if new_para_text:
    #                 new_para_text += " " + sentence_text
    #             else:
    #                 new_para_text = sentence_text

    #         para.text = new_para_text

    #     doc.save(output_doc_path)
    #     logging.info(f"Document saved with in-text citations: {output_doc_path}")


    def process_document(self, input_doc_path: str, output_doc_path: str):
        """
        Processes the input document, performs vector search within the limited
        collection, and saves the updated document with in-text citations.
        """
        logging.info(f"Processing document: {input_doc_path}")
        doc = Document(input_doc_path)

        for paragraph_index, para in enumerate(doc.paragraphs, start=1):
            if not para.text.strip():
                continue

            logging.info(f"Processing paragraph {paragraph_index}.")
            sentences = list(nlp(para.text).sents)
            new_para_text = ""

            for sentence_index, sent in enumerate(sentences, start=1):
                sentence_text = sent.text.strip()
                if not sentence_text:
                    continue

                # Check if the sentence is a standalone heading
                if self.is_heading(sentence_text):
                    logging.info(f"Skipping heading: {sentence_text}")
                    new_para_text += f" {sentence_text}"  # Preserve headings in the output
                    continue

                logging.info(f"Processing sentence {sentence_index}: {sentence_text}")
                relevant_papers = self.find_relevant_papers(sentence_text)

                if relevant_papers:
                    citation_texts = []
                    for paper_doc in relevant_papers:
                        paper_metadata = paper_doc.metadata
                        title = paper_metadata.get("title")

                        if title:
                            citation_texts.append(f"({title})")

                    if citation_texts:
                        combined_citation = " ".join(citation_texts)
                        sentence_text += f" {combined_citation}"

                if new_para_text:
                    new_para_text += " " + sentence_text
                else:
                    new_para_text = sentence_text

            para.text = new_para_text

        doc.save(output_doc_path)
        logging.info(f"Document saved with in-text citations: {output_doc_path}")

# if __name__ == "__main__":
#     COLLECTION_NAME = "quantum_physics"  # Replace with your collection name

#     inserter = SemanticCitationInserter(
#         mongo_uri=MONGODB_ATLAS_CLUSTER_URI,
#         db_name=MONGO_DB_NAME,
#         collection_name=COLLECTION_NAME
#     )

#     input_doc = "/Users/naija/Documents/gigs/tweakr/tweakr-mvp/test_docs/testdoc.docx"
#     output_doc = "/Users/naija/Documents/gigs/tweakr/tweakr-mvp/test_docs/testdoc_with_citations.docx"

#     inserter.process_document(input_doc, output_doc)
