# import os
# from docx import Document
# from sqlalchemy.orm import Session
# from typing import List, Optional
# from datetime import datetime
# from app.core.references_generator import ReferenceGenerator
# from datapipeline.core.database import get_session_with_ctx_manager
# from datapipeline.core.mongo_client import MongoDBVectorStoreManager
# from datapipeline.models.papers import Papers
# from datapipeline.core.constants import MONGODB_ATLAS_CLUSTER_URI, MONGO_DB_NAME
# from datapipeline.core.utils import embeddings



# import os
# import spacy
# import numpy as np
# from docx import Document
# from typing import List, Dict, Tuple
# from sqlalchemy.orm import Session
# from dataclasses import dataclass

# # Assuming you've installed spacy model:
# # python -m spacy download en_core_web_sm
# nlp = spacy.load("en_core_web_sm")


# @dataclass
# class PaperMetadata:
#     title: str
#     authors: List[str]
#     year: str
#     category: str


# class PaperMetadataRetriever:
#     """
#     Retrieves full paper metadata from Postgres given a title.
#     """
#     def get_paper_metadata(self, title: str) -> Optional[PaperMetadata]:
#         with get_session_with_ctx_manager() as session:
#             paper = session.query(Papers).filter(Papers.title == title).first()
#             if not paper:
#                 return None
            
#             # Parse authors
#             authors = []
#             if paper.authors:
#                 try:
#                     if paper.authors.strip().startswith("["):
#                         authors = json.loads(paper.authors)
#                     else:
#                         authors = [a.strip() for a in paper.authors.split(",")]
#                 except:
#                     authors = [a.strip() for a in paper.authors.split(",")]

#             year = "n.d."
#             if paper.pub_date and isinstance(paper.pub_date, (date, datetime)):
#                 year = str(paper.pub_date.year)
            
#             return PaperMetadataDataclass(
#                 title=paper.title,
#                 authors=authors,
#                 year=year,
#                 category=paper.category
#             )


# class InTextCitationFormatter:
#     """
#     Formats in-text citations according to APA style for this example.
#     """
#     def format_citation(self, paper_meta: PaperMetadata) -> str:
#         # APA in-text: (AuthorLastName, Year)
#         # If multiple authors: (Author1 et al., Year)
#         authors = paper_meta.authors
#         year = paper_meta.year
#         if not authors:
#             return f"(Unknown, {year})"

#         author_lastnames = [a.split()[-1] for a in authors]
#         if len(author_lastnames) == 1:
#             citation = f"({author_lastnames[0]}, {year})"
#         elif len(author_lastnames) == 2:
#             citation = f"({author_lastnames[0]} & {author_lastnames[1]}, {year})"
#         else:
#             citation = f"({author_lastnames[0]} et al., {year})"
#         return citation



# class SemanticCitationInserter:
#     """
#     Uses a vector store to find semantically related papers for each sentence in a document
#     and inserts in-text citations at the end of the sentence.
#     """
#     def __init__(self,
#         collection_name: str, 
#         threshold: float = 0.75,
#         top_k: int = 3
#     ):
        
#         self.threshold = threshold
#         self.top_k = top_k
#         self.collection_name = collection_name
        
#         # Embedding model for sentences
#         self.embedding_model = embeddings
        
#         # Initialize MongoDB vector store
#         # self.mongo_client = MongoClient(mongo_uri)
#         # db = self.mongo_client[db_name]
#         # collection = db[collection_name]
        
#         # self.vector_store = MongoDBAtlasVectorSearch(
#         #     collection=collection,
#         #     embedding=self.embedding_model,
#         #     index_name=f"{collection_name}_index"
#         # )

#         self.manager = MongoDBVectorStoreManager(
#             connection_string=MONGODB_ATLAS_CLUSTER_URI, 
#             db_name=MONGO_DB_NAME
#         )

#         self.citation_formatter = InTextCitationFormatter()
#         self.paper_retriever = PaperMetadataRetriever()

#     def embed_sentence(self, sentence: str) -> np.ndarray:
#         return np.array(self.embedding_model.embed_query(sentence), dtype=float)

#     def find_relevant_papers(self, sentence_embedding: np.ndarray):
#         """
#         Query the vectorstore with the sentence embedding to find top_k relevant papers.
#         """
#         embedding_list = sentence_embedding.tolist()
#         # results = self.vector_store.similarity_search_by_vector(embedding_list, k=self.top_k)
#         results = self.manager.similarity_search_by_vector(self.collection_name, embedding_list, k=self.top_k)
#         return results

#     def process_document(self, input_doc_path: str, output_doc_path: str):
#         doc = Document(input_doc_path)

#         for para in doc.paragraphs:
#             sentences = list(nlp(para.text).sents)
#             new_para_text = ""
            
#             for sent in sentences:
#                 sentence_text = sent.text.strip()
#                 if not sentence_text:
#                     continue

#                 # Embed the sentence
#                 sentence_embedding = self.embed_sentence(sentence_text)
                
#                 # Find relevant papers
#                 relevant_paper_docs = self.find_relevant_papers(sentence_embedding)

#                 citation_texts = []
#                 for paper_doc in relevant_paper_docs:
#                     paper_title = paper_doc.metadata.get("title")
#                     if paper_title:
#                         paper_meta = self.paper_retriever.get_paper_metadata(paper_title)
#                         if paper_meta:
#                             citation_text = self.citation_formatter.format_citation(paper_meta)
#                             citation_texts.append(citation_text)

#                 if citation_texts:
#                     # Append citations to the sentence
#                     combined_citation = " ".join(citation_texts)
#                     sentence_text = f"{sentence_text} {combined_citation}"

#                 # Reconstruct paragraph text
#                 if new_para_text:
#                     new_para_text += " " + sentence_text
#                 else:
#                     new_para_text = sentence_text

#             para.text = new_para_text

#         doc.save(output_doc_path)
#         print(f"Document saved with in-text citations at {output_doc_path}")



# if __name__ == "__main__":
#     COLLECTION_NAME = "quantum_physics"  # The collection where you stored vector embeddings

#     inserter = SemanticCitationInserter(
#         collection_name=COLLECTION_NAME,
#         threshold=0.75
#     )

#     input_doc = "/Users/naija/Documents/gigs/tweakr/tweakr-mvp/test_docs/testdoc.docx"
#     output_doc = "/Users/naija/Documents/gigs/tweakr/tweakr-mvp/test_docs/testdoc_with_citations.docx"
#     inserter.process_document(input_doc, output_doc)


import os
import logging
from docx import Document
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
from app.core.references_generator import ReferenceGenerator
from datapipeline.core.database import get_session_with_ctx_manager
from datapipeline.core.mongo_client import MongoDBVectorStoreManager
from datapipeline.models.papers import Papers
from datapipeline.core.constants import MONGODB_ATLAS_CLUSTER_URI, MONGO_DB_NAME
from datapipeline.core.utils import embeddings
import spacy
import numpy as np
from dataclasses import dataclass

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Load SpaCy model
nlp = spacy.load("en_core_web_sm")

@dataclass
class PaperMetadata:
    title: str
    authors: List[str]
    year: str
    category: str

class PaperMetadataRetriever:
    def get_paper_metadata(self, title: str) -> Optional[PaperMetadata]:
        with get_session_with_ctx_manager() as session:
            paper = session.query(Papers).filter(Papers.title == title).first()
            if not paper:
                return None

            authors = []
            if paper.authors:
                try:
                    if paper.authors.strip().startswith("["):
                        authors = json.loads(paper.authors)
                    else:
                        authors = [a.strip() for a in paper.authors.split(",")]
                except:
                    authors = [a.strip() for a in paper.authors.split(",")]

            year = "n.d."
            if paper.pub_date and isinstance(paper.pub_date, (datetime, str)):
                year = str(paper.pub_date.year)

            return PaperMetadata(
                title=paper.title,
                authors=authors,
                year=year,
                category=paper.category
            )

class InTextCitationFormatter:
    def format_citation(self, paper_meta: PaperMetadata) -> str:
        authors = paper_meta.authors
        year = paper_meta.year
        if not authors:
            return f"(Unknown, {year})"

        author_lastnames = [a.split()[-1] for a in authors]
        if len(author_lastnames) == 1:
            return f"({author_lastnames[0]}, {year})"
        elif len(author_lastnames) == 2:
            return f"({author_lastnames[0]} & {author_lastnames[1]}, {year})"
        else:
            return f"({author_lastnames[0]} et al., {year})"

class SemanticCitationInserter:
    def __init__(self, collection_name: str, threshold: float = 0.75, top_k: int = 3):
        self.threshold = threshold
        self.top_k = top_k
        self.collection_name = collection_name
        self.embedding_model = embeddings
        self.manager = MongoDBVectorStoreManager(
            connection_string=MONGODB_ATLAS_CLUSTER_URI, 
            db_name=MONGO_DB_NAME
        )
        self.citation_formatter = InTextCitationFormatter()
        self.paper_retriever = PaperMetadataRetriever()

    def embed_sentence(self, sentence: str) -> np.ndarray:
        logger.info(f"Embedding sentence: {sentence}")
        return np.array(self.embedding_model.embed_query(sentence), dtype=float)

    def find_relevant_papers(self, sentence_embedding: np.ndarray):
        logger.info(f"Searching for relevant papers using embedding.")
        embedding_list = sentence_embedding.tolist()
        results = self.manager.similarity_search_by_vector(self.collection_name, embedding_list, k=self.top_k)
        logger.info(f"Found {len(results)} relevant papers.")
        return results

    def process_document(self, input_doc_path: str, output_doc_path: str):
        logger.info(f"Processing document: {input_doc_path}")
        doc = Document(input_doc_path)

        for para_idx, para in enumerate(doc.paragraphs):
            logger.info(f"Processing paragraph {para_idx + 1}.")
            sentences = list(nlp(para.text).sents)
            new_para_text = ""

            for sent_idx, sent in enumerate(sentences):
                sentence_text = sent.text.strip()
                if not sentence_text:
                    continue

                logger.info(f"Processing sentence {sent_idx + 1}: {sentence_text}")

                # Embed the sentence
                sentence_embedding = self.embed_sentence(sentence_text)

                # Find relevant papers
                relevant_paper_docs = self.find_relevant_papers(sentence_embedding)

                citation_texts = []
                for paper_doc in relevant_paper_docs:
                    paper_title = paper_doc.metadata.get("title")
                    if paper_title:
                        paper_meta = self.paper_retriever.get_paper_metadata(paper_title)
                        if paper_meta:
                            citation_text = self.citation_formatter.format_citation(paper_meta)
                            citation_texts.append(citation_text)

                if citation_texts:
                    combined_citation = " ".join(citation_texts)
                    sentence_text = f"{sentence_text} {combined_citation}"
                    logger.info(f"Appended citation(s): {combined_citation}")

                new_para_text += (" " if new_para_text else "") + sentence_text

            para.text = new_para_text

        doc.save(output_doc_path)
        logger.info(f"Document saved with in-text citations: {output_doc_path}")

if __name__ == "__main__":
    COLLECTION_NAME = "quantum_physics"

    inserter = SemanticCitationInserter(
        collection_name=COLLECTION_NAME,
        threshold=0.75
    )

    input_doc = "/Users/naija/Documents/gigs/tweakr/tweakr-mvp/test_docs/testdoc.docx"
    output_doc = "/Users/naija/Documents/gigs/tweakr/tweakr-mvp/test_docs/testdoc_with_citations.docx"
    inserter.process_document(input_doc, output_doc)
