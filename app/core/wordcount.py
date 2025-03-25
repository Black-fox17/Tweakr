import docx
import re
import os

def count_words_in_docx(file_path):
    """
    Count words in a Word document.
    
    Args:
        file_path (str): Path to the Word document
    
    Returns:
        dict: Dictionary containing word count statistics
    """
    try:
        # Check if file exists
        if not os.path.exists(file_path):
            return {"error": f"File not found: {file_path}"}
        
        # Check if file is a Word document
        if not file_path.endswith(('.docx', '.doc')):
            return {"error": f"Not a Word document: {file_path}"}
        
        # Load the document
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Join all paragraphs with space
        text = ' '.join(full_text)
        
        # Count words (any sequence of non-whitespace characters)
        words = re.findall(r'\S+', text)
        num_words = len(words)
        
        # Count characters
        num_chars = len(text)
        num_chars_no_spaces = len(text.replace(" ", ""))
        
        # Count paragraphs
        num_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        
        return {
            "word_count": num_words,
            "character_count": num_chars,
            "character_count_no_spaces": num_chars_no_spaces,
            "paragraph_count": num_paragraphs,
        }
    
    except Exception as e:
        return {"error": f"Error processing document: {str(e)}"}
