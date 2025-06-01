import os
import json
import uuid
import random
import logging
import spacy
from docx import Document
from typing import List, Dict, Any, Optional
import requests
import time
from urllib.parse import quote_plus
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# Setup logging
logging.basicConfig(level=logging.INFO)

class AcademicCitationProcessor:
    """
    Enhanced replacement for MongoDB-based citation processor using multiple academic search APIs.
    Provides fallback mechanisms and improved error handling with optimized API usage.
    """
    
    def __init__(self, style="APA", search_providers=None, threshold=0.0, top_k=3, max_api_calls=500):
        """
        Initialize the academic citation processor.

        Parameters:
        - style (str): Citation style (APA, MLA, Chicago).
        - search_providers (list): List of academic search providers in priority order.
        - threshold (float): Minimum relevance threshold.
        - top_k (int): Number of top relevant documents to retrieve per provider.
        - max_api_calls (int): Maximum number of API calls allowed (increased default to 500).
        """
        self.style = style
        self.search_providers = search_providers or ["semantic_scholar", "crossref", "openalex"]
        self.threshold = threshold
        self.top_k = top_k
        self.max_api_calls = max_api_calls
        self.api_call_count = 0
        self.matched_paper_titles = []
        
        # Enhanced caching for better API efficiency
        self.search_cache = {}
        self.cache_hits = 0
        
        # Load SpaCy model for sentence segmentation
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logging.error("SpaCy model 'en_core_web_sm' not found. Please install it with: python -m spacy download en_core_web_sm")
            raise
        
        # Initialize session with retry strategy
        self.session = requests.Session()
        retry_strategy = Retry(
            total=3,
            status_forcelist=[429, 500, 502, 503, 504],
            backoff_factor=1
        )
        adapter = HTTPAdapter(max_retries=retry_strategy)
        self.session.mount("http://", adapter)
        self.session.mount("https://", adapter)
        
        # Optimized rate limiting delays for each provider (seconds)
        self.rate_limits = {
            'semantic_scholar': 0.05,  # Reduced from 0.1
            'crossref': 0.03,          # Reduced from 0.05
            'openalex': 0.03           # Reduced from 0.05
        }
        
        # Track last API call time for each provider
        self.last_api_call = {}
        
        # Academic domains for filtering relevance
        self.academic_domains = [
            'healthcare', 'medicine', 'biology', 'computer science', 'machine learning',
            'business', 'management', 'marketing', 'mathematics', 'physics', 
            'neuroscience', 'psychology', 'education', 'engineering', 'nursing',
            'public health', 'clinical', 'research', 'evidence-based', 'systematic'
        ]

    def _respect_rate_limit(self, provider: str):
        """Ensure rate limiting is respected for each provider"""
        if provider in self.last_api_call:
            time_since_last = time.time() - self.last_api_call[provider]
            required_delay = self.rate_limits.get(provider, 0.03)
            if time_since_last < required_delay:
                time.sleep(required_delay - time_since_last)
        
        self.last_api_call[provider] = time.time()

    def _get_cache_key(self, query: str, provider: str) -> str:
        """Generate cache key for search queries"""
        return f"{provider}:{query.lower().strip()}"

    def is_dynamic_heading(self, para) -> bool:
        """
        Dynamically detects if a paragraph is a heading or subheading using multiple heuristics.
        """
        text = para.text.strip()
        if not text:
            return False

        try:
            style_name = para.style.name.lower()
            if "heading" in style_name or "title" in style_name:
                logging.info(f"Detected heading based on style: '{text}'")
                return True
        except Exception:
            pass

        # Heuristic: short text without punctuation or bullet points
        words = text.split()
        if len(words) < 8 and not any(punct in text for punct in [".", "?", "!", ";", ":"]):
            # Skip bullet points and numbered lists
            if not (text.startswith('-') or text.startswith('•') or text[0].isdigit()):
                logging.info(f"Detected potential heading based on text heuristic: '{text}'")
                return True

        return False

    def clean_query(self, query: str) -> str:
        """Clean and optimize search query"""
        if not query or not isinstance(query, str):
            return ""
        
        # Remove bullet points and numbering
        query = query.strip()
        if query.startswith('-') or query.startswith('•'):
            query = query[1:].strip()
        
        # Remove leading numbers
        if query and query[0].isdigit() and '.' in query[:5]:
            query = query.split('.', 1)[1].strip()
        
        # Limit query length but be more generous
        words = query.split()
        if len(words) > 20:  # Increased from 15
            query = ' '.join(words[:20])
        
        return query

    def _estimate_api_calls_needed(self, sentences: List[str]) -> int:
        """Estimate how many API calls will be needed for processing sentences"""
        unique_queries = set()
        for sentence in sentences:
            cleaned = self.clean_query(sentence)
            if cleaned:
                # Check if already cached
                cache_key_base = cleaned.lower().strip()
                is_cached = any(cache_key_base in key for key in self.search_cache.keys())
                if not is_cached:
                    unique_queries.add(cleaned)
        
        # Estimate calls per unique query (number of providers)
        return len(unique_queries) * len(self.search_providers)

    def _select_sentences_intelligently(self, all_sentences: List[str], max_calls_available: int) -> List[str]:
        """
        Intelligently select sentences based on available API calls and content quality.
        Prioritizes longer, more substantial sentences and ensures randomization.
        """
        if not all_sentences:
            return []
        
        # Filter out very short sentences (less likely to need citations)
        substantial_sentences = [
            s for s in all_sentences 
            if len(s.strip()) >= 20 and len(s.split()) >= 4
        ]
        
        if not substantial_sentences:
            substantial_sentences = all_sentences
        
        # Estimate API calls needed per sentence (average across providers)
        calls_per_sentence = len(self.search_providers)
        max_sentences = max(1, max_calls_available // calls_per_sentence)
        
        if len(substantial_sentences) <= max_sentences:
            return substantial_sentences
        
        # Prioritize sentences by length and academic keywords
        def sentence_priority(sentence):
            score = len(sentence.split()) * 0.1  # Length factor
            # Boost sentences with academic keywords
            for domain in self.academic_domains:
                if domain.lower() in sentence.lower():
                    score += 1
            # Boost sentences with numbers/statistics
            if any(char.isdigit() for char in sentence):
                score += 0.5
            return score
        
        # Sort by priority but add randomization
        sentence_scores = [(s, sentence_priority(s) + random.random() * 0.5) for s in substantial_sentences]
        sentence_scores.sort(key=lambda x: x[1], reverse=True)
        
        selected = [s[0] for s in sentence_scores[:max_sentences]]
        
        # Add some randomization to the final selection
        if len(selected) > 2:
            # Keep top performers but randomize the rest
            top_portion = selected[:len(selected)//2]
            remaining = selected[len(selected)//2:]
            random.shuffle(remaining)
            selected = top_portion + remaining
        
        return selected

    def search_all_providers(self, query: str, max_results: int = None) -> List[Dict]:
        """
        Search for academic papers using all available providers with fallback and caching.
        
        Parameters:
        - query (str): Search query text
        - max_results (int): Maximum number of results to return per provider
        
        Returns:
        - List[Dict]: Combined list of paper metadata from all providers
        """
        if self.api_call_count >= self.max_api_calls:
            logging.warning(f"Maximum API calls ({self.max_api_calls}) exceeded")
            return []
        
        query = self.clean_query(query)
        if not query:
            return []
        
        max_results = max_results or self.top_k
        all_papers = []
        seen_titles = set()
        
        # Check cache first
        cache_results = {}
        for provider in self.search_providers:
            cache_key = self._get_cache_key(query, provider)
            if cache_key in self.search_cache:
                cache_results[provider] = self.search_cache[cache_key]
                self.cache_hits += 1
        
        for provider in self.search_providers:
            if self.api_call_count >= self.max_api_calls:
                break
                
            try:
                # Use cached results if available
                if provider in cache_results:
                    papers = cache_results[provider]
                    logging.info(f"Using cached results for {provider}: {len(papers)} papers")
                else:
                    logging.info(f"Searching {provider} for: '{query[:50]}...'")
                    papers = self._search_provider(provider, query, max_results)
                    
                    # Cache the results
                    cache_key = self._get_cache_key(query, provider)
                    self.search_cache[cache_key] = papers
                
                # Deduplicate by title
                unique_papers = []
                for paper in papers:
                    title = paper.get('title', '').lower().strip()
                    if title and title not in seen_titles:
                        seen_titles.add(title)
                        unique_papers.append(paper)
                
                all_papers.extend(unique_papers)
                logging.info(f"Found {len(unique_papers)} unique papers from {provider}")
                
                # If we have enough papers, we can stop searching other providers
                if len(all_papers) >= max_results * 2:
                    logging.info(f"Sufficient papers found ({len(all_papers)}), stopping search")
                    break
                    
            except Exception as e:
                logging.error(f"Error searching {provider}: {e}")
                continue
        
        return all_papers

    def _search_provider(self, provider: str, query: str, max_results: int) -> List[Dict]:
        """Search a specific provider"""
        self._respect_rate_limit(provider)
        self.api_call_count += 1
        
        if provider == 'semantic_scholar':
            return self._search_semantic_scholar(query, max_results)
        elif provider == 'crossref':
            return self._search_crossref(query, max_results)
        elif provider == 'openalex':
            return self._search_openalex(query, max_results)
        else:
            logging.error(f"Unsupported search provider: {provider}")
            return []

    def _search_semantic_scholar(self, query: str, max_results: int) -> List[Dict]:
        """Search using Semantic Scholar API"""
        try:
            url = 'https://api.semanticscholar.org/graph/v1/paper/search'
            params = {
                'query': query,
                'limit': min(max_results, 50),  # Reduced from 100 for efficiency
                'fields': 'title,authors,abstract,year,venue,citationCount,url,paperId'
            }
            headers = {'Accept': 'application/json'}
            
            response = self.session.get(url, params=params, headers=headers, timeout=10)
            response.raise_for_status()
            
            data = response.json()
            papers = []
            
            for paper in data.get('data', []):
                # Safely extract paper data
                title = paper.get('title')
                if not title:
                    continue
                
                authors = []
                for author in paper.get('authors', []):
                    if author and isinstance(author, dict):
                        name = author.get('name')
                        if name:
                            authors.append(name)
                
                paper_data = {
                    'title': title,
                    'authors': authors,
                    'abstract': paper.get('abstract') or '',
                    'year': paper.get('year'),
                    'venue': paper.get('venue') or '',
                    'url': paper.get('url') or '',
                    'citations': paper.get('citationCount', 0),
                    'paper_id': paper.get('paperId') or '',
                    'source': 'Semantic Scholar'
                }
                papers.append(paper_data)
            
            logging.info(f"Retrieved {len(papers)} papers from Semantic Scholar")
            return papers
            
        except requests.exceptions.RequestException as e:
            if hasattr(e, 'response') and e.response.status_code == 429:
                logging.warning("Semantic Scholar rate limit hit, backing off...")
                time.sleep(2)
            raise e
        except Exception as e:
            logging.error(f"Error searching Semantic Scholar: {e}")
            return []

    def _search_crossref(self, query: str, max_results: int) -> List[Dict]:
        """Search using Crossref API"""
        try:
            url = 'https://api.crossref.org/works'
            params = {
                'query': query,
                'rows': min(max_results, 50),  # Reduced from 100
                'sort': 'relevance'
            }
            headers = {'Accept': 'application/json'}
            
            response = self.session.get(url, params=params, headers=headers, timeout=10)
            response.raise_for_status()
            
            data = response.json()
            papers = []
            
            for item in data.get('message', {}).get('items', []):
                # Extract title
                title_list = item.get('title', [])
                if not title_list:
                    continue
                title = ' '.join(title_list) if isinstance(title_list, list) else str(title_list)
                
                # Extract authors
                authors = []
                for author in item.get('author', []):
                    if author and isinstance(author, dict):
                        if 'given' in author and 'family' in author:
                            authors.append(f"{author['given']} {author['family']}")
                        elif 'family' in author:
                            authors.append(author['family'])
                
                # Extract year
                year = None
                try:
                    if 'published-print' in item:
                        year = item['published-print']['date-parts'][0][0]
                    elif 'published-online' in item:
                        year = item['published-online']['date-parts'][0][0]
                except (IndexError, KeyError, TypeError):
                    pass
                
                # Extract venue
                venue = ''
                container_title = item.get('container-title', [])
                if container_title:
                    venue = container_title[0] if isinstance(container_title, list) else str(container_title)
                
                paper_data = {
                    'title': title,
                    'authors': authors,
                    'abstract': item.get('abstract', ''),
                    'year': year,
                    'venue': venue,
                    'url': item.get('URL', ''),
                    'citations': item.get('is-referenced-by-count', 0),
                    'doi': item.get('DOI', ''),
                    'source': 'Crossref'
                }
                papers.append(paper_data)
            
            logging.info(f"Retrieved {len(papers)} papers from Crossref")
            return papers
            
        except Exception as e:
            logging.error(f"Error searching Crossref: {e}")
            return []

    def _search_openalex(self, query: str, max_results: int) -> List[Dict]:
        """Search using OpenAlex API"""
        try:
            url = 'https://api.openalex.org/works'
            params = {
                'search': query,
                'per-page': min(max_results, 50),  # Reduced from 100
                'sort': 'relevance_score:desc'
            }
            headers = {'Accept': 'application/json'}
            
            response = self.session.get(url, params=params, headers=headers, timeout=10)
            response.raise_for_status()
            
            data = response.json()
            papers = []
            
            for work in data.get('results', []):
                title = work.get('title')
                if not title:
                    continue
                
                # Extract authors
                authors = []
                for authorship in work.get('authorships', []):
                    if authorship and isinstance(authorship, dict):
                        author = authorship.get('author', {})
                        if author and author.get('display_name'):
                            authors.append(author['display_name'])
                
                # Extract venue
                venue = ''
                primary_location = work.get('primary_location', {})
                if primary_location and isinstance(primary_location, dict):
                    source = primary_location.get('source', {})
                    if source and isinstance(source, dict):
                        venue = source.get('display_name', '')
                
                # Extract URL
                url_val = ''
                if primary_location and isinstance(primary_location, dict):
                    url_val = primary_location.get('landing_page_url', '')
                
                paper_data = {
                    'title': title,
                    'authors': authors,
                    'abstract': work.get('abstract', ''),
                    'year': work.get('publication_year'),
                    'venue': venue,
                    'url': url_val,
                    'citations': work.get('cited_by_count', 0),
                    'doi': work.get('doi', ''),
                    'source': 'OpenAlex'
                }
                papers.append(paper_data)
            
            logging.info(f"Retrieved {len(papers)} papers from OpenAlex")
            return papers
            
        except Exception as e:
            logging.error(f"Error searching OpenAlex: {e}")
            return []

    def calculate_relevance_score(self, sentence: str, paper: Dict) -> float:
        """
        Calculate a relevance score between a sentence and a paper.
        Enhanced scoring system with domain relevance and author quality check.
        """
        if not sentence or not isinstance(sentence, str):
            return 0.0
        
        # Check if paper has valid authors - this is a requirement
        authors = paper.get('authors', [])
        if not authors or len(authors) == 0:
            return 0.0
        
        # Check for valid author names
        valid_authors = [author for author in authors if author and str(author).strip()]
        if not valid_authors:
            return 0.0
        
        sentence_lower = sentence.lower()
        title = paper.get('title', '')
        abstract = paper.get('abstract', '')
        
        # Handle None values safely
        if not title:
            title = ''
        if not abstract:
            abstract = ''
        
        title_lower = title.lower()
        abstract_lower = abstract.lower()
        
        # Simple keyword matching score
        sentence_words = set(sentence_lower.split())
        title_words = set(title_lower.split())
        abstract_words = set(abstract_lower.split())
        
        # Remove common stop words for better matching
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should'}
        sentence_words = sentence_words - stop_words
        title_words = title_words - stop_words
        abstract_words = abstract_words - stop_words
        
        if not sentence_words:  # If no meaningful words left
            return 0.0
        
        # Calculate overlap scores
        title_overlap = len(sentence_words.intersection(title_words)) / len(sentence_words)
        abstract_overlap = len(sentence_words.intersection(abstract_words)) / len(sentence_words)
        
        # Weight title matches higher than abstract matches
        relevance_score = (title_overlap * 0.8) + (abstract_overlap * 0.2)
        
        # Check for academic domain relevance
        domain_boost = 0
        for domain in self.academic_domains:
            if domain in sentence_lower or domain in title_lower or domain in abstract_lower:
                domain_boost += 0.1
        
        relevance_score += min(domain_boost, 0.2)  # Cap domain boost
        
        # Boost score for recent papers and high citation counts
        year = paper.get('year')
        if year and isinstance(year, (int, str)) and str(year).isdigit():
            year_int = int(year)
            if year_int >= 2020:
                relevance_score *= 1.2
            elif year_int >= 2015:
                relevance_score *= 1.1
            elif year_int >= 2010:
                relevance_score *= 1.05
        
        citations = paper.get('citations', 0)
        if isinstance(citations, (int, str)) and str(citations).isdigit():
            citations_int = int(citations)
            if citations_int > 100:
                relevance_score *= 1.1
            elif citations_int > 50:
                relevance_score *= 1.05
            elif citations_int > 10:
                relevance_score *= 1.02
        
        # Boost papers with multiple authors (indicates collaborative research)
        if len(valid_authors) > 1:
            relevance_score *= 1.05
        
        # Boost papers from reputable venues
        venue = paper.get('venue', '').lower()
        high_impact_keywords = ['journal', 'proceedings', 'conference', 'review', 'nature', 'science', 'ieee', 'acm']
        if any(keyword in venue for keyword in high_impact_keywords):
            relevance_score *= 1.05
        
        return min(relevance_score, 1.0)  # Cap at 1.0

    def find_relevant_papers(self, sentence: str, return_all: bool = False) -> List[Dict]:
        """
        Find relevant papers for a sentence using multiple academic search APIs with fallback.
        Returns only the best citation per sentence and excludes papers without authors.
        
        Parameters:
        - sentence (str): The sentence to find relevant papers for
        - return_all (bool): If True, returns all papers above threshold; 
                            If False, returns only the best matching paper
        
        Returns:
        - List[Dict]: List of relevant papers with metadata (max 1 paper per sentence)
        """
        if not isinstance(sentence, str) or not sentence.strip():
            logging.warning(f"Skipping empty or invalid sentence: '{sentence}'")
            return []

        # Search for papers across all providers
        papers = self.search_all_providers(sentence)
        if not papers:
            return []

        # Filter papers and calculate relevance scores
        relevant_papers = []
        for paper in papers:
            try:
                # Skip papers without authors
                authors = paper.get('authors', [])
                if not authors or len(authors) == 0:
                    logging.debug(f"Skipping paper without authors: '{paper.get('title', 'Unknown')}'")
                    continue
                
                # Skip if authors list contains only empty strings or None values
                valid_authors = [author for author in authors if author and str(author).strip()]
                if not valid_authors:
                    logging.debug(f"Skipping paper with invalid authors: '{paper.get('title', 'Unknown')}'")
                    continue
                
                # Update paper with valid authors only
                paper['authors'] = valid_authors
                
                # Calculate relevance score
                relevance_score = self.calculate_relevance_score(sentence, paper)
                if relevance_score >= self.threshold:
                    paper['relevance_score'] = relevance_score
                    relevant_papers.append(paper)
                    
            except Exception as e:
                logging.error(f"Error calculating relevance for paper '{paper.get('title', 'Unknown')}': {e}")
                continue

        # Sort by relevance score (best first)
        relevant_papers.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
        
        if return_all:
            return relevant_papers
        else:
            # Return only the best paper (top 1) per sentence
            return relevant_papers[:1] if relevant_papers else []

    def format_citation(self, authors: List[str], year: str) -> str:
        """
        Format in-text citation based on the specified style.
        """
        if not authors:
            authors = ["Unknown"]
        
        # Handle year safely
        if not year or year == 'None':
            year = "n.d."
        else:
            year = str(year)
        
        # Extract first name from the first author
        first_author = authors[0] if authors else "Unknown"
        if " " in first_author:
            first_name = first_author.split(" ")[0]
        else:
            first_name = first_author

        if self.style == "APA":
            if len(authors) == 1:
                return f"({first_name}, {year})"
            else:
                return f"({first_name} et al., {year})"
        
        elif self.style == "MLA":
            if len(authors) == 1:
                return f"({first_name} {year})"
            else:
                return f"({first_name} et al. {year})"
        
        elif self.style == "Chicago":
            if len(authors) == 1:
                return f"({first_name} {year})"
            else:
                return f"({first_name} et al. {year})"
        
        else:
            raise ValueError(f"Unsupported citation style: {self.style}")

    def prepare_citations_for_review(self, input_path: str, max_paragraphs: int = 200, 
                                   random_sample: bool = True, strict_api_limit: bool = True) -> Dict[str, Any]:
        """
        Prepare in-text citations for frontend review using multiple academic search APIs.
        Enhanced version with better error handling, intelligent sentence selection, and API optimization.

        Parameters:
        - input_path (str): Path to the input document.
        - max_paragraphs (int): Maximum number of paragraphs to process (increased default to 200).
        - random_sample (bool): Whether to randomly sample paragraphs (default: True).
        - strict_api_limit (bool): Whether to strictly enforce API limits with intelligent selection.

        Returns:
        - Dict containing document-level and citation-level information for review.
        """
        logging.info(f"Preparing citations for review using providers {self.search_providers} from file: '{input_path}'")
        logging.info(f"Max API calls available: {self.max_api_calls}")

        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file '{input_path}' does not exist.")

        try:
            doc = Document(input_path)
            total_paragraphs = len(doc.paragraphs)
            logging.info(f"Document paragraphs count: {total_paragraphs}")
            
            # Calculate paragraphs to process
            paragraphs_to_process = min(total_paragraphs, max_paragraphs)
            logging.info(f"Will process {paragraphs_to_process} out of {total_paragraphs} paragraphs")
            
            # Select paragraphs
            if random_sample and total_paragraphs > paragraphs_to_process:
                paragraph_indices = sorted(random.sample(range(total_paragraphs), paragraphs_to_process))
                paragraphs_to_process_list = [doc.paragraphs[i] for i in paragraph_indices]
            else:
                paragraphs_to_process_list = doc.paragraphs[:paragraphs_to_process]
                paragraph_indices = list(range(paragraphs_to_process))
            
            # Pre-collect all sentences for intelligent selection
            all_sentences = []
            sentence_metadata = []
            
            for para_idx, para in enumerate(paragraphs_to_process_list):
                actual_para_idx = paragraph_indices[para_idx] + 1 if random_sample else para_idx + 1
                paragraph_text = para.text.strip()

                # Skip empty paragraphs and headings
                if not paragraph_text or self.is_dynamic_heading(para):
                    continue

                # Tokenize paragraph into sentences
                try:
                    sentences = list(self.nlp(paragraph_text).sents)
                except Exception as tokenize_error:
                    logging.error(f"Error tokenizing paragraph {para_idx}: {tokenize_error}")
                    continue

                for sent_idx, sent in enumerate(sentences, start=1):
                    sentence_text = sent.text.strip()
                    if sentence_text and len(sentence_text) >= 10:
                        all_sentences.append(sentence_text)
                        sentence_metadata.append({
                            'paragraph_idx': para_idx,
                            'actual_para_idx': actual_para_idx,
                            'sentence_idx': sent_idx,
                            'original_document_index': paragraph_indices[para_idx] if random_sample else para_idx
                        })

            logging.info(f"Total sentences extracted: {len(all_sentences)}")
            
            # Intelligent sentence selection based on API limits
            if strict_api_limit and len(all_sentences) > 0:
                api_calls_available = self.max_api_calls - self.api_call_count
                logging.info(f"API calls available: {api_calls_available}")
                
                if api_calls_available > 0:
                    selected_sentences = self._select_sentences_intelligently(all_sentences, api_calls_available)
                    logging.info(f"Intelligently selected {len(selected_sentences)} sentences for processing")
                    
                    # Update metadata to match selected sentences
                    selected_indices = []
                    for selected in selected_sentences:
                        try:
                            selected_index = all_sentences.index(selected)
                            selected_indices.append(selected_index)
                        except ValueError:
                            logging.warning(f"Selected sentence not found in all_sentences: '{selected}'")
                    
                    # Update metadata to match selected sentences
                    for idx, selected_idx in enumerate(selected_indices):
                        sentence_metadata[selected_idx]['selected_sentence'] = all_sentences[selected_idx]
                    
                    # Update API call count
                    self.api_call_count += len(selected_indices)
                    logging.info(f"Updated API call count: {self.api_call_count}")

            # Process remaining sentences if any
            if len(all_sentences) - self.api_call_count > 0:
                remaining_sentences = all_sentences[self.api_call_count:]
                logging.info(f"Processing {len(remaining_sentences)} remaining sentences")

                # Process remaining sentences
                for sentence in remaining_sentences:
                    try:
                        # Find relevant papers
                        papers = self.find_relevant_papers(sentence)
                        if papers:
                            # Format citation
                            citation = self.format_citation(papers[0]['authors'], papers[0]['year'])
                            logging.info(f"Formatted citation: {citation}")
                        else:
                            logging.warning(f"No relevant papers found for sentence: '{sentence}'")
                    except Exception as e:
                        logging.error(f"Error processing sentence '{sentence}': {e}")
                        continue


    
