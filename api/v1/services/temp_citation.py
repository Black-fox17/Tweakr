import os
import json
import uuid
import random
import logging
import spacy
from docx import Document
from typing import List, Dict, Any
import asyncio
import aiohttp
from collections import defaultdict
from async_lru import alru_cache
from scholarly import scholarly, ProxyGenerator
from app.core.gemini_helper import get_document_context_with_gemini


# Setup logging
logging.basicConfig(level=logging.INFO)

class TempCitationProcessor:
    """
    Enhanced replacement for MongoDB-based citation processor using multiple academic search APIs.
    Provides fallback mechanisms and improved error handling with proper termination controls.
    """
    
    def __init__(self, style="APA", search_providers=None, threshold=0.0, top_k=5, max_api_calls=None, additional_context=""):
        self.style = style
        self.search_providers = search_providers or ["google_scholar", "semantic_scholar", "crossref", "openalex"]
        self.threshold = threshold
        self.top_k = top_k
        self.max_api_calls = max_api_calls
        self.api_call_count = 0
        self.matched_paper_titles = []
        self.additional_context = additional_context
        self.research_context = ""
        self.document_category = ""
        self.field_keywords = []
        
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logging.error("SpaCy model 'en_core_web_sm' not found. Please install it with: python -m spacy download en_core_web_sm")
            raise
        
        # Configure scholarly to use a proxy if needed, to avoid getting blocked
        pg = ProxyGenerator()
        # Consider using a free proxy or a premium one if you have it
        # success = pg.FreeProxies()
        # scholarly.use_proxy(pg)

    def _calculate_api_limits_and_eta(self, total_sentences: int) -> tuple:
        if self.max_api_calls is not None:
            return self.max_api_calls, 0
        
        citation_rate = 1
        avg_providers_per_search = len(self.search_providers)
        estimated_citations = int(total_sentences * citation_rate)
        calculated_max_calls = min(estimated_citations * avg_providers_per_search, 1000)
        
        avg_time_per_call = 0.2  # Reduced avg time due to async nature
        estimated_eta_seconds = (calculated_max_calls / avg_providers_per_search) * avg_time_per_call
        
        self.max_api_calls = int(calculated_max_calls)
        
        return self.max_api_calls, estimated_eta_seconds

    def smart_sentence_selection(self, all_sentences: list, max_sentences: int = None) -> list:
        if not all_sentences:
            return []
        
        if max_sentences == None:
            max_sentences = min(len(all_sentences), 200)
        
        if len(all_sentences) <= max_sentences:
            return all_sentences
        
        # Simplified selection: prioritize sentences with academic keywords
        academic_keywords = {
            'study', 'research', 'analysis', 'data', 'results', 'findings', 'evidence',
            'method', 'approach', 'theory', 'model', 'framework', 'hypothesis',
            'significant', 'correlation', 'impact', 'effect', 'relationship',
            'according', 'reported', 'demonstrated', 'showed', 'indicated'
        }
        
        priority_sentences = [s for s in all_sentences if any(kw in s['text'].lower() for kw in academic_keywords)]
        regular_sentences = [s for s in all_sentences if s not in priority_sentences]

        selected = []
        if len(priority_sentences) >= max_sentences:
            selected = random.sample(priority_sentences, max_sentences)
        else:
            selected.extend(priority_sentences)
            needed = max_sentences - len(selected)
            selected.extend(random.sample(regular_sentences, min(needed, len(regular_sentences))))
            
        return selected

    def is_dynamic_heading(self, para) -> bool:
        text = para.text.strip()
        if not text:
            return False
        try:
            if "heading" in para.style.name.lower():
                return True
        except Exception:
            pass
        words = text.split()
        return len(words) < 8 and not any(punct in text for punct in ".?!;:")
    
    def enhance_query_with_context(self, original_query: str, sentence_context: str = "") -> str:
        enhanced_query = original_query
        if self.document_category:
            enhanced_query = f"{enhanced_query} {self.document_category}"
        
        if self.field_keywords:
            field_terms = " ".join(self.field_keywords[:3])
            enhanced_query = f"{enhanced_query} {field_terms}"
        
        return self.clean_query(enhanced_query)

    def clean_query(self, query: str) -> str:
        query = query.strip()
        if query.startswith('-') or query.startswith('•'):
            query = query[1:].strip()
        if query and query[0].isdigit() and '.' in query[:5]:
            query = query.split('.', 1)[1].strip()
        words = query.split()
        return ' '.join(words[:15])

    @alru_cache(maxsize=1024)
    async def search_all_providers_async(self, query: str, max_results: int = None) -> List[Dict]:
        if self.api_call_count >= self.max_api_calls:
            return []
        
        query = self.clean_query(query)
        if not query:
            return []
        
        max_results = max_results or self.top_k
        
        connector = aiohttp.TCPConnector(limit=20, limit_per_host=10)
        timeout = aiohttp.ClientTimeout(total=10, connect=4)
        
        async with aiohttp.ClientSession(connector=connector, timeout=timeout) as session:
            tasks = []
            enhanced_query = self.enhance_query_with_context(query, self.research_context)
            print(enhanced_query)
            for provider in self.search_providers:
                if self.api_call_count >= self.max_api_calls:
                    break
                
                self.api_call_count += 1
                task = self._search_provider_async(session, provider, query, max_results)
                tasks.append(task)
            
            if not tasks:
                return []
            
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            all_papers = []
            seen_titles = set()
            
            for result in results:
                if isinstance(result, Exception):
                    logging.warning(f"Search provider call failed: {result}")
                    continue
                
                for paper in result:
                    title = paper.get('title', '').lower().strip()
                    if title and title not in seen_titles:
                        seen_titles.add(title)
                        all_papers.append(paper)
        
        return all_papers
    
    async def _search_provider_async(self, session: aiohttp.ClientSession, provider: str, query: str, max_results: int) -> List[Dict]:
        try:
            if provider == 'google_scholar':
                return await self._search_google_scholar_async(query, max_results)
            elif provider == 'semantic_scholar':
                return await self._search_semantic_scholar_async(session, query, max_results)
            elif provider == 'crossref':
                return await self._search_crossref_async(session, query, max_results)
            elif provider == 'openalex':
                return await self._search_openalex_async(session, query, max_results)
        except Exception as e:
            logging.error(f"Failed to search {provider}: {e}")
            return []

    async def _search_google_scholar_async(self, query: str, max_results: int) -> List[Dict]:
        try:
            loop = asyncio.get_running_loop()
            # scholarly.search_pubs is not async, so we run it in an executor
            search_results = await loop.run_in_executor(None, lambda: scholarly.search_pubs(query))
            
            papers = []
            for i, pub in enumerate(search_results):
                if i >= max_results:
                    break
                papers.append({
                    'title': pub['bib'].get('title'),
                    'authors': pub['bib'].get('author'),
                    'year': pub['bib'].get('pub_year'),
                    'venue': pub['bib'].get('venue'),
                    'url': pub.get('pub_url'),
                    'citations': pub.get('num_citations', 0),
                    'source': 'Google Scholar',
                })
            return papers
        except Exception as e:
            logging.error(f"Error searching Google Scholar: {e}")
            return []

    async def _search_semantic_scholar_async(self, session: aiohttp.ClientSession, query: str, max_results: int) -> List[Dict]:
        url = 'https://api.semanticscholar.org/graph/v1/paper/search'
        params = {'query': query, 'limit': max_results, 'fields': 'title,authors,year,venue,citationCount,url'}
        try:
            async with session.get(url, params=params) as response:
                if response.status == 429: await asyncio.sleep(1); return []
                response.raise_for_status()
                data = await response.json()
                return [{
                    'title': p.get('title'), 'authors': [a.get('name') for a in p.get('authors', [])],
                    'year': p.get('year'), 'venue': p.get('venue'), 'url': p.get('url'),
                    'citations': p.get('citationCount', 0), 'source': 'Semantic Scholar'
                } for p in data.get('data', []) if p.get('title') and p.get('authors')]
        except Exception as e:
            logging.error(f"Error searching Semantic Scholar: {e}"); return []

    async def _search_crossref_async(self, session: aiohttp.ClientSession, query: str, max_results: int) -> List[Dict]:
        url = 'https://api.crossref.org/works'
        params = {'query': query, 'rows': max_results, 'sort': 'relevance'}
        try:
            async with session.get(url, params=params) as response:
                response.raise_for_status()
                data = await response.json()
                papers = []
                for item in data.get('message', {}).get('items', []):
                    authors = [f"{a.get('given', '')} {a.get('family', '')}".strip() for a in item.get('author', [])]
                    if not item.get('title') or not authors: continue
                    year = item.get('published-print', {}).get('date-parts', [[None]])[0][0]
                    papers.append({
                        'title': ' '.join(item.get('title', [])), 'authors': authors, 'year': year,
                        'venue': ' '.join(item.get('container-title', [])), 'url': item.get('URL'),
                        'citations': item.get('is-referenced-by-count', 0), 'source': 'Crossref'
                    })
                return papers
        except Exception as e:
            logging.error(f"Error searching Crossref: {e}"); return []

    async def _search_openalex_async(self, session: aiohttp.ClientSession, query: str, max_results: int) -> List[Dict]:
        url = 'https://api.openalex.org/works'
        params = {'search': query, 'per-page': max_results, 'sort': 'relevance_score:desc'}
        try:
            async with session.get(url, params=params) as response:
                response.raise_for_status()
                data = await response.json()
                papers = []
                for work in data.get('results', []):
                    authors = [a['author'].get('display_name') for a in work.get('authorships', [])]
                    if not work.get('title') or not authors: continue
                    venue = work.get('primary_location', {}).get('source', {}).get('display_name')
                    papers.append({
                        'title': work.get('title'), 'authors': authors, 'year': work.get('publication_year'),
                        'venue': venue, 'url': work.get('primary_location', {}).get('landing_page_url'),
                        'citations': work.get('cited_by_count', 0), 'source': 'OpenAlex'
                    })
                return papers
        except Exception as e:
            logging.error(f"Error searching OpenAlex: {e}"); return []
        
    def calculate_relevance_score(self, sentence: str, paper: Dict) -> float:
        if not sentence or not isinstance(sentence, str) or not paper.get('authors'):
            return 0.0
        
        sentence_lower = sentence.lower()
        title = (paper.get('title') or '').lower()
        
        stop_words = {'the', 'a', 'an', 'and', 'or', 'in', 'on', 'to', 'for', 'of', 'is', 'are', 'was', 'were'}
        sentence_words = set(sentence_lower.split()) - stop_words
        title_words = set(title.split()) - stop_words
        
        if not sentence_words: return 0.0
        
        title_overlap = len(sentence_words.intersection(title_words)) / len(sentence_words)
        score = title_overlap * 0.8
        
        year = paper.get('year')
        if year and str(year).isdigit():
            if int(year) >= 2020: score *= 1.2
            elif int(year) >= 2015: score *= 1.1
        
        citations = paper.get('citations', 0)
        if citations > 100: score *= 1.1
        elif citations > 50: score *= 1.05
        
        return min(score, 1.0)

    async def batch_process_sentences_async(self, sentences: list) -> list:
        tasks = [self.process_single_sentence_async(s) for s in sentences if self.api_call_count < self.max_api_calls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        all_citations = []
        for res in results:
            if isinstance(res, Exception):
                logging.error(f"Error processing sentence: {res}")
            elif res:
                all_citations.append(res)
        return all_citations

    async def process_single_sentence_async(self, sentence_data: dict) -> dict:
        try:
            sentence_text = sentence_data['text']
            papers = await self.search_all_providers_async(sentence_text)
            if not papers: return None
            
            relevant_papers = []
            for paper in papers:
                score = self.calculate_relevance_score(sentence_text, paper)
                if score >= self.threshold:
                    paper['relevance_score'] = score
                    relevant_papers.append(paper)
            
            if not relevant_papers: return None
            
            best_paper = max(relevant_papers, key=lambda x: x.get('relevance_score', 0))
            year = best_paper.get('year')
            if not year or (str(year).isdigit() and int(year) < 2015):
                return None

            return {
                "id": str(uuid.uuid4()),
                "original_sentence": sentence_text,
                "paper_details": {
                    "title": best_paper.get('title'), "authors": best_paper.get('authors', []),
                    "year": str(year) if year else "n.d.", "url": best_paper.get('url', ''),
                    "doi": best_paper.get('doi', ''), "venue": best_paper.get('venue', ''),
                    "citations": best_paper.get('citations', 0), "relevance_score": round(best_paper.get('relevance_score', 0), 3),
                    "source": best_paper.get('source', 'Unknown'),
                },
                "metadata": {
                    "paragraph_index": sentence_data['actual_para_idx'],
                    "sentence_index": sentence_data['sent_idx'],
                }
            }
        except Exception as e:
            logging.error(f"Error in process_single_sentence_async: {e}")
            return None

    async def prepare_citations_for_review(self, input_path: str, max_paragraphs: int = 100) -> Dict[str, Any]:
        logging.info(f"Preparing citations for review from file: '{input_path}'")
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file '{input_path}' does not exist.")

        doc = Document(input_path)
        
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Get context from Gemini
        context_data = await get_document_context_with_gemini(full_text, self.additional_context)
        print(context_data)
        self.research_context = context_data.get("research_context", "")
        self.document_category = context_data.get("document_category", "")
        self.field_keywords = context_data.get("field_keywords", [])
        
        logging.info(f"Gemini context acquired: Category='{self.document_category}', Context='{self.research_context}'")

        paragraphs_to_process = doc.paragraphs[:min(len(doc.paragraphs), max_paragraphs)]
        
        all_sentences = []
        for para_idx, para in enumerate(paragraphs_to_process):
            if not para.text.strip() or self.is_dynamic_heading(para): continue
            try:
                for sent_idx, sent in enumerate(self.nlp(para.text.strip()).sents, 1):
                    if len(sent.text.strip()) >= 15:
                        all_sentences.append({
                            'text': sent.text.strip(), 'actual_para_idx': para_idx + 1, 'sent_idx': sent_idx
                        })
            except Exception as e:
                logging.error(f"Error tokenizing paragraph {para_idx}: {e}")

        total_sentences = len(all_sentences)
        calculated_max_calls, estimated_eta = self._calculate_api_limits_and_eta(total_sentences)
        selected_sentences = self.smart_sentence_selection(all_sentences, min(total_sentences, 150))
        
        citations = await self.batch_process_sentences_async(selected_sentences)
        
        return {
            "document_id": str(uuid.uuid4()),
            "total_citations": len(citations),
            "citations": citations,
            "context_info": {
                "research_context": self.research_context,
                "document_category": self.document_category,
                "field_keywords": self.field_keywords,
                "detected_domain": self.document_category # Use the Gemini-detected category
            },
            "diagnostics": {
                "processed_paragraphs": len(set(s['actual_para_idx'] for s in selected_sentences)),
                "processed_sentences": len(selected_sentences),
                "api_calls_made": self.api_call_count,
                "max_api_calls": calculated_max_calls,
                "estimated_eta_seconds": estimated_eta,
            }
        }